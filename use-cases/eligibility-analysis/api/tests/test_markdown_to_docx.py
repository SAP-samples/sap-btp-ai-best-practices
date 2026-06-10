"""Tests for the markdown-to-DOCX converter."""

import tempfile
import unittest
from pathlib import Path

from app.optimizer.report.markdown_to_docx import markdown_to_docx


_SAMPLE_MARKDOWN = """\
# Test Report

**Generated:** 2025-01-28

---

## Selection Metrics

| Metric | Value |
|--------|-------|
| Candidates | 100 |
| Selected | 80 |

## Details

- **Item A**: Description
- Item B: Another

### Sub-section

1. First
2. Second

Paragraph with **bold** inline.
"""


class TestMarkdownToDocx(unittest.TestCase):
    def test_creates_valid_docx(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "test.docx"
            result = markdown_to_docx(_SAMPLE_MARKDOWN, output)

            self.assertEqual(result, output)
            self.assertTrue(output.exists())
            self.assertGreater(output.stat().st_size, 1000)

            # DOCX is a ZIP file (PK header)
            with open(output, "rb") as f:
                header = f.read(2)
            self.assertEqual(header, b"PK")

    def test_empty_markdown(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "empty.docx"
            result = markdown_to_docx("", output)
            self.assertTrue(output.exists())

    def test_table_rendered(self) -> None:
        """Tables should produce a docx table object."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "table.docx"
            markdown_to_docx(_SAMPLE_MARKDOWN, output)

            doc = Document(str(output))
            self.assertTrue(len(doc.tables) > 0, "Expected at least one table in DOCX")
            # First table should have 3 rows (header + 2 data)
            self.assertEqual(len(doc.tables[0].rows), 3)


if __name__ == "__main__":
    unittest.main()
