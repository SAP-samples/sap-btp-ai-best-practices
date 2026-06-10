"""Generic markdown-to-DOCX converter using *python-docx*.

Parses the same subset of markdown that ``markdown_to_pdf`` handles (headers,
tables, bold, bullets, numbered lists, horizontal rules) and produces a Word
document with styling that matches the PDF output (teal headers, Open Sans
font, matching table structure).
"""

from __future__ import annotations

import re
import logging
import zipfile
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


def _write_fallback_docx(markdown_content: str, output_path: Path) -> Path:
    """Create a minimal DOCX-compatible ZIP when python-docx is unavailable."""
    escaped = (
        markdown_content.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    paragraphs = []
    for line in markdown_content.splitlines() or [""]:
        text = line or " "
        paragraphs.append(
            "<w:p><w:r><w:t xml:space=\"preserve\">"
            + (
                text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            + "</w:t></w:r></w:p>"
        )

    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:body>"
        + "".join(paragraphs)
        + "</w:body></w:document>"
    )

    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/fallback.txt", escaped)
    logger.warning("python-docx not installed; wrote fallback DOCX to %s", output_path)
    return output_path


def _add_runs_with_inline_styles(paragraph, text: str) -> None:
    """Render inline markdown emphasis (``**bold**`` and ``*italic*``)."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue

        run = None
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)


def _set_cell_text_with_inline_styles(cell, text: str, force_bold: bool = False) -> None:
    """Replace cell content and render inline markdown emphasis."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _add_runs_with_inline_styles(paragraph, text)
    if force_bold:
        for run in paragraph.runs:
            run.bold = True


def _apply_heading_style(paragraph, level: int) -> None:
    """Apply teal/dark-gray heading styling that mirrors the PDF look."""
    from docx.shared import Pt, RGBColor

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 134, 171)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 134, 171)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(70, 70, 70)


def _render_markdown_table_docx(doc, table_lines: List[str]) -> None:
    """Parse markdown table lines and render a DOCX table."""
    if len(table_lines) < 2:
        return

    header = [cell.strip() for cell in table_lines[0].split("|") if cell.strip()]
    if not header:
        return

    data_rows: List[List[str]] = []
    for line in table_lines[2:]:
        if line.strip().startswith("|"):
            cells = [cell.strip() for cell in line.split("|") if cell.strip()]
            if cells:
                while len(cells) < len(header):
                    cells.append("")
                data_rows.append(cells[:len(header)])

    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = "Table Grid"

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        _set_cell_text_with_inline_styles(cell, text, force_bold=True)

    for i, row in enumerate(data_rows, start=1):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            _set_cell_text_with_inline_styles(cell, text)

    doc.add_paragraph("")


def markdown_to_docx(markdown_content: str, output_path: Path) -> Path:
    """Convert markdown text to a styled DOCX document.

    Args:
        markdown_content: Markdown string to render.
        output_path: Destination file path (will be created/overwritten).

    Returns:
        The *output_path* for convenience.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ModuleNotFoundError:
        return _write_fallback_docx(markdown_content, output_path)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Open Sans"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(50, 50, 50)

    lines = markdown_content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Table block
        if line.strip().startswith("|"):
            table_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _render_markdown_table_docx(doc, table_lines)
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(line[2:])
            _apply_heading_style(p, 1)
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:])
            _apply_heading_style(p, 2)
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:])
            _apply_heading_style(p, 3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_inline_styles(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_inline_styles(p, line)
        elif line.strip() == "---":
            doc.add_paragraph("")
        elif "**" in line or "*" in line:
            p = doc.add_paragraph()
            _add_runs_with_inline_styles(p, line)
        elif line.strip():
            doc.add_paragraph(line)

        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("DOCX report written to %s", output_path)
    return output_path
